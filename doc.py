#!/usr/bin/env python3
# pip install python-docx
# document object -> entire word document
# Document object contains list of paragraph objects
# Each paragraph object contains list of run objects
import docx
d = docx.Document('demo.docx')
print(d.paragraphs) # list of paragraph objects
print(d.paragraphs[0].text) # Document Title
print(d.paragraphs[1].text) # A plain paragraph having some bold and some italic.
p = d.paragraphs[1]
print(p.runs) # list of run objects
print(p.runs[0].text) # each run occurs when there is a change in styling
print(p.runs[3].italic) # returns True if it's italic, otherwise returns None
print(p.runs[3].bold) # returns True if it's bold, otherwise returns None
p.runs[3].text = 'italic and underlined'
d.save('demo2.docx') # save the original doc as demo2.docx


def createDocx(docName):
    d = docx.Document()
    d.add_paragraph('Hello this is a paragraph.')
    d.add_paragraph('Hello this is another paragraph.')
    p  = d.paragraphs[0]
    p.add_run(' This is a new run.') # it adds another run; so two runs now
    p.runs[1].bold = True # bolds the second run
    d.save(docName)

def readDoc(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    return '\n'.join(fullText)


print(readDoc('demo3.docx'))

# createDocx('demo3.docx')

